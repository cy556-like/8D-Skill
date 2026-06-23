#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
8D ?????????
=====================
?????????????????,?? 8D ?? .xlsx(? Sheet)? .docx ???

??:
    - openpyxl(?? Excel)
    - python-docx(?? Word)

?????? pip install?

??:
    python3 generate_8d.py \\
        --product "??????(???)" \\
        --defect "????/??" \\
        --customer "XX??????" \\
        --defect-rate "500PPM" \\
        --batch-size "15" \\
        --template paint-defect \\
        --output-dir ~/Desktop

?? ??????(?? SKILL.md ???):
    - defect-rate ????? 500 PPM (0.05%),???? 50 PPM;
      ???? 3% / 5% / 8% / 11.5% ???????????
    - batch-size ???????? 8D ???????/???????,
      ?? 1-30 ?;??????,???? 500 / 2000 ????????
"""

import argparse
import json
import os
import subprocess
import sys
import datetime
from pathlib import Path

# ============================================================
# ?????????
# ============================================================

REQUIRED_PACKAGES = {
    "openpyxl": "openpyxl",
    "docx": "python-docx",  # ??:???? docx,??? python-docx
}


def ensure_packages():
    """?????????? Python ??"""
    import importlib

    for import_name, pip_name in REQUIRED_PACKAGES.items():
        try:
            importlib.import_module(import_name)
        except ImportError:
            print(f"[INFO] {pip_name} ???,??????...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name, "--quiet"]
                )
                print(f"[INFO] {pip_name} ?????")
            except subprocess.CalledProcessError as e:
                print(f"[ERROR] {pip_name} ????:{e}")
                print(f"        ?????:pip install {pip_name}")
                sys.exit(1)


ensure_packages()

# ????????
import openpyxl
from openpyxl.styles import (
    Alignment,
    Border,
    Side,
    PatternFill,
    Font,
)
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================
# ??
# ============================================================

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent  # 8d-report/
TEMPLATES_DIR = SKILL_ROOT / "templates"

# ??
HEADER_FILL = "003366"  # ?????
ALT_ROW_FILL = "D6E4F0"  # ?????
ROOT_CAUSE_FILL = "FFF2CC"  # ?????????
SUBHEADER_FILL = "4472C4"  # ????
YELLOW_SECTION_FILL = "FFF8E1"  # ????????

# ??
FONT_NAME = "????"
FONT_NAME_EN = "Microsoft YaHei"


# ============================================================
# ????
# ============================================================

def replace_placeholders(text, context):
    """??????????"""
    if not isinstance(text, str):
        return text
    return (
        text.replace("{defect_type}", context.get("defect", ""))
        .replace("{product_name}", context.get("product", ""))
        .replace("{customer}", context.get("customer", ""))
    )


def replace_placeholders_deep(obj, context):
    """??????/????????"""
    if isinstance(obj, str):
        return replace_placeholders(obj, context)
    if isinstance(obj, list):
        return [replace_placeholders_deep(item, context) for item in obj]
    if isinstance(obj, dict):
        return {k: replace_placeholders_deep(v, context) for k, v in obj.items()}
    return obj


def generate_8d_number():
    """?? 8D ??:8D-YYYYMMDD-HHMMSS"""
    now = datetime.datetime.now()
    return now.strftime("8D-%Y%m%d-%H%M%S")


def safe_filename(name):
    """?????????????(??????)?"""
    illegal_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '\n', '\r', '\t']
    for ch in illegal_chars:
        name = name.replace(ch, '_')
    # ????????
    if len(name) > 50:
        name = name[:50]
    return name


def load_template(template_slug, context):
    """???? JSON ???????"""
    template_path = TEMPLATES_DIR / template_slug / "template.json"
    if not template_path.exists():
        print(f"[WARN] ?? {template_slug} ???,??? generic-defect")
        template_path = TEMPLATES_DIR / "generic-defect" / "template.json"

    with open(template_path, "r", encoding="utf-8") as f:
        template = json.load(f)

    # ???????
    template = replace_placeholders_deep(template, context)
    return template


# ============================================================
# Excel ??
# ============================================================

def get_thin_border():
    """??????"""
    side = Side(border_style="thin", color="000000")
    return Border(left=side, right=side, top=side, bottom=side)


def get_header_font():
    return Font(name=FONT_NAME, size=11, bold=True, color="FFFFFF")


def get_body_font():
    return Font(name=FONT_NAME, size=10, color="000000")


def get_header_fill():
    return PatternFill(start_color=HEADER_FILL, end_color=HEADER_FILL, fill_type="solid")


def get_alt_fill():
    return PatternFill(start_color=ALT_ROW_FILL, end_color=ALT_ROW_FILL, fill_type="solid")


def get_root_cause_fill():
    return PatternFill(start_color=ROOT_CAUSE_FILL, end_color=ROOT_CAUSE_FILL, fill_type="solid")


def get_subheader_fill():
    return PatternFill(start_color=SUBHEADER_FILL, end_color=SUBHEADER_FILL, fill_type="solid")


def get_subheader_font():
    return Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")


def get_yellow_section_fill():
    """????????????"""
    return PatternFill(start_color=YELLOW_SECTION_FILL, end_color=YELLOW_SECTION_FILL, fill_type="solid")


def get_yellow_section_font():
    """??????????(???)?"""
    return Font(name=FONT_NAME, size=12, bold=True, color=HEADER_FILL)


def apply_yellow_section_style(cell):
    """???????????"""
    cell.font = get_yellow_section_font()
    cell.fill = get_yellow_section_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_header_style(cell):
    """???????"""
    cell.font = get_header_font()
    cell.fill = get_header_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_body_style(cell, row_idx_in_data, is_root_cause=False):
    """??????(??? / ????)?"""
    cell.font = get_body_font()
    if is_root_cause:
        cell.fill = get_root_cause_fill()
    elif row_idx_in_data % 2 == 1:
        cell.fill = get_alt_fill()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def apply_subheader_style(cell):
    """?????????"""
    cell.font = get_subheader_font()
    cell.fill = get_subheader_fill()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = get_thin_border()


def set_column_widths(ws, widths):
    """?????"""
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_table(ws, start_row, headers, rows, col_widths, root_cause_row_indices=None):
    """
    ????????????(??? + ???)?
    root_cause_row_indices: ???????????????(0-based ?????)
    ??:??????
    """
    if root_cause_row_indices is None:
        root_cause_row_indices = []

    # ??
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        apply_header_style(cell)
    ws.row_dimensions[start_row].height = 28

    # ???
    for i, row_data in enumerate(rows):
        current_row = start_row + 1 + i
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=value)
            is_rc = i in root_cause_row_indices
            apply_body_style(cell, i, is_root_cause=is_rc)
        # ????(????????)
        max_len = max((len(str(v)) for v in row_data), default=10)
        ws.row_dimensions[current_row].height = max(20, min(80, 15 + max_len // 4))

    # ????
    set_column_widths(ws, col_widths)

    # ????
    ws.freeze_panes = ws.cell(row=start_row + 1, column=1)

    return start_row + 1 + len(rows)


def write_section_title(ws, row, title, span_cols=6):
    """??????(?????,?????)?"""
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span_cols)
    cell = ws.cell(row=row, column=1, value=title)
    apply_header_style(cell)
    ws.row_dimensions[row].height = 30
    return row + 1


def write_yellow_section_title(ws, row, title, span_cols=6):
    """????????(??????,??????????)?"""
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span_cols)
    cell = ws.cell(row=row, column=1, value=title)
    apply_yellow_section_style(cell)
    ws.row_dimensions[row].height = 30
    return row + 1


def write_kv_block(ws, start_row, kv_pairs, label_col_width=22, value_col_width=40):
    """
    ??????(??? + ??),2 ????
    kv_pairs: [(label, value), ...]
    """
    for i, (label, value) in enumerate(kv_pairs):
        current_row = start_row + i
        # ???
        label_cell = ws.cell(row=current_row, column=1, value=label)
        label_cell.font = Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")
        label_cell.fill = get_subheader_fill()
        label_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        label_cell.border = get_thin_border()
        # ??
        value_cell = ws.cell(row=current_row, column=2, value=value)
        value_cell.font = get_body_font()
        value_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        value_cell.border = get_thin_border()
        ws.row_dimensions[current_row].height = max(20, min(60, 15 + len(str(value)) // 6))

    set_column_widths(ws, [label_col_width, value_col_width])
    return start_row + len(kv_pairs)


# ============================================================
# Excel ??(? Sheet ??:D0-D8 ???????)
# ============================================================

def generate_excel(context, template, output_path, report_number):
    """?? Excel ??(? Sheet:8D??)?"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "8D??"
    ws.sheet_properties.tabColor = "003366"

    row = 1
    d0_d2 = template.get("d0_d2", {})

    # ==================== D0-D2 ====================
    row = write_yellow_section_title(ws, row, "D0-D2  ???????????", span_cols=4)
    row += 1

    # ?????
    row = write_kv_block(ws, row, [
        ("8D ????", report_number),
        ("??????", datetime.datetime.now().strftime("%Y-%m-%d")),
        ("?????", "____"),
        ("????", "???"),
    ])
    row += 1

    # ????
    row = write_section_title(ws, row, "??????", span_cols=2)
    row = write_kv_block(ws, row, [
        ("????", context.get("customer", "____")),
        ("?????", "____"),
        ("??????", "____"),
        ("??????", "____"),
        ("??????", "____(?? / ?? / 8D??? / ??)"),
    ])
    row += 1

    # ????
    row = write_section_title(ws, row, "??????", span_cols=2)
    row = write_kv_block(ws, row, [
        ("????", context.get("product", "____")),
        ("???? / ???", "____"),
        ("???", "____"),
        ("????", context.get("batch_size", "____")),
        ("????", "____"),
        ("???", context.get("defect_rate", "____")),
        ("????", d0_d2.get("defect_level_hint", "____")),
        ("????", d0_d2.get("discovery_location_hint", "____")),
        ("??", d0_d2.get("impact_hint", "____")),
    ])
    row += 1

    # D2 5W2H
    row = write_section_title(ws, row, "??D2 ????(5W2H)", span_cols=2)
    row = write_kv_block(ws, row, [
        ("What(????)", context.get("defect", "____")),
        ("When(????)", "____(?? / ??)"),
        ("Where(????)", "____(?? / ???? / ????)"),
        ("Who(????)", "____"),
        ("Why(??????)", "____(????? / ????)"),
        ("How(????)", "____(???? / ????)"),
        ("How many(????/?)", f"{context.get('defect_rate', '____')},?? {context.get('batch_size', '____')}"),
    ])
    row += 1

    # ????
    row = write_section_title(ws, row, "??????", span_cols=2)
    problem_stmt = (
        f"???{context.get('product', '____')}?? {d0_d2.get('discovery_location_hint', '____')} "
        f"???{context.get('defect', '____')}???,"
        f"??? {context.get('defect_rate', '____')},"
        f"?????? {context.get('batch_size', '____')}?"
        f"?????? {d0_d2.get('impact_hint', '____')},?????? 8D ???????"
    )
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    cell = ws.cell(row=row, column=1, value=problem_stmt)
    cell.font = get_body_font()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    cell.border = get_thin_border()
    ws.row_dimensions[row].height = 60
    row += 1

    # D1 ??
    row += 1
    row = write_section_title(ws, row, "??D1 ?????", span_cols=4)
    row = write_table(ws, row,
        ["??", "??", "??", "????"],
        [
            ["????(?????)", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["????", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["SQE(????????)", "____", "???", "____"],
            ["????(??)", "____", "???", "____"],
        ],
        col_widths=[28, 18, 14, 22],
    )

    # ==================== D3 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D3  ??????(Interim Containment Actions)", span_cols=6)
    row += 1
    row = write_section_title(ws, row, "????????", span_cols=6)
    d3 = template.get("d3_template", {})
    actions = d3.get("containment_actions", [])
    while len(actions) < 5:
        actions.append("____(???????)")
    row = write_table(ws, row,
        ["??", "????", "???", "????", "????", "??"],
        [[i + 1, actions[i], "____", "____", "____(?100%????/?????)", "???"] for i in range(5)],
        col_widths=[6, 50, 12, 14, 28, 10],
    )
    row += 1
    row = write_section_title(ws, row, "?????????", span_cols=2)
    row = write_kv_block(ws, row, [
        ("??????", context.get("defect_rate", "____")),
        ("??????", "____"),
        ("??????(24h?)", "____"),
        ("??????(72h?)", "____"),
        ("????", "____(?? / ???,????????)"),
        ("????????", "____"),
        ("??????", "____"),
    ])

    # ==================== D4 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D4  ??????(Root Cause Analysis)", span_cols=4)
    row += 1
    row = write_section_title(ws, row, "??5Why ??", span_cols=4)
    d4 = template.get("d4_template", {})
    five_why = d4.get("5why_path", {})
    steps = five_why.get("steps", [])
    why_rows = []
    root_cause_indices = []

    if steps:
        for idx, step in enumerate(steps):
            level = step.get("level", "____")
            question = step.get("question", "____")
            answer = step.get("answer", "____")
            evidence = step.get("evidence", "____")
            why_rows.append([level, question, answer, evidence])
            if "??" in str(level):
                root_cause_indices.append(idx)
    else:
        problem = five_why.get("problem", "____")
        why_rows.append(["??", problem, "-", "-"])
        why1 = five_why.get("why1", "____")
        why_rows.append(["Why 1", "?????????", why1, "____"])
        why2_hints = five_why.get("why2_hints", [])
        why_rows.append(["Why 2", f"???:{why1[:30]}...", "\n".join(f". {h}" for h in why2_hints) if why2_hints else "____", "____"])
        why3_hints = five_why.get("why3_hints", [])
        why_rows.append(["Why 3", "???????????", "\n".join(f". {h}" for h in why3_hints) if why3_hints else "____", "____"])
        why4_hints = five_why.get("why4_hints", [])
        why_rows.append(["Why 4", "?????/???????", "\n".join(f". {h}" for h in why4_hints) if why4_hints else "____", "____"])
        why5_root = five_why.get("why5_root", "____")
        why_rows.append(["Why 5(??)", "?????/???????", why5_root, "____"])
        root_cause_indices = [5]

    row = write_table(ws, row,
        ["??", "??", "??", "??"],
        why_rows,
        col_widths=[14, 38, 50, 22],
        root_cause_row_indices=root_cause_indices,
    )

    row += 1
    row = write_section_title(ws, row, "????? 6M ??", span_cols=4)
    six_m = d4.get("6m_analysis", {})

    if isinstance(six_m, list):
        m_rows = []
        m_rc_indices = []
        for idx, item in enumerate(six_m):
            m_name = item.get("m", "____")
            finding = item.get("finding", "____")
            judgment = item.get("judgment", "____")
            m_rows.append([m_name, finding, judgment])
            if "????" in str(judgment):
                m_rc_indices.append(idx)
        row = write_table(ws, row,
            ["6M ??", "????", "??"],
            m_rows,
            col_widths=[18, 55, 18],
            root_cause_row_indices=m_rc_indices,
        )
    else:
        row = write_table(ws, row,
            ["6M ??", "????", "??", "????"],
            [
                ["Man(?)", six_m.get("man", "____"), "____", "____(?/?)"],
                ["Machine(?)", six_m.get("machine", "____"), "____", "____(?/?)"],
                ["Material(?)", six_m.get("material", "____"), "____", "____(?/?)"],
                ["Method(?)", six_m.get("method", "____"), "____", "____(?/?)"],
                ["Measurement(?)", six_m.get("measurement", "____"), "____", "____(?/?)"],
                ["Environment(?)", six_m.get("environment", "____"), "____", "____(?/?)"],
            ],
            col_widths=[16, 50, 24, 14],
        )

    row += 1
    row = write_section_title(ws, row, "????????", span_cols=2)
    rc_summary = d4.get("root_cause_summary", [])
    if rc_summary:
        rc_kv = [(f"{rc.get('id', '____')}({rc.get('type', '____')})", rc.get("description", "____")) for rc in rc_summary]
        row = write_kv_block(ws, row, rc_kv, label_col_width=22, value_col_width=70)
    else:
        row = write_kv_block(ws, row, [
            ("RC1(????)", "____(?? 5Why ? 1-2 ???)"),
            ("RC2(????)", "____(?? 5Why ? 3-4 ???)"),
            ("RC3(????)", "____(?? 5Why ? 5 ???)"),
        ], label_col_width=22, value_col_width=70)

    row += 1
    row = write_section_title(ws, row, "??????????", span_cols=2)
    verify_text = d4.get("verification", "")
    if verify_text:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        cell = ws.cell(row=row, column=1, value=verify_text)
        cell.font = get_body_font()
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        cell.border = get_thin_border()
        ws.row_dimensions[row].height = max(30, min(120, 15 + len(verify_text) // 4))
    else:
        row = write_kv_block(ws, row, [
            ("????", "____(???? / ???? / ????)"),
            ("????", "____"),
            ("????", "____(??? / ??????)"),
            ("??? / ??", "____"),
        ], label_col_width=22, value_col_width=70)

    # ==================== D5-D6 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D5-D6  ??????(Permanent Corrective Actions)", span_cols=7)
    row += 1
    row = write_section_title(ws, row, "??D5 CA ??????", span_cols=6)
    d5_d6 = template.get("d5_d6_template", {})
    permanent_actions = d5_d6.get("permanent_actions", [])

    ca_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            ca_rows.append([i, action.get("action", "____"), action.get("target", "____"), "____(?/?/?)", "____(???????)", "____(??/??)"])
        else:
            ca_rows.append([i, str(action), "____", "____(?/?/?)", "____(???????)", "____(??/??)"])
    while len(ca_rows) < 3:
        ca_rows.append([len(ca_rows) + 1, "____(??? CA ??)", "____", "____", "____", "____"])

    row = write_table(ws, row,
        ["??", "CA ??", "????", "???", "????", "??"],
        ca_rows,
        col_widths=[6, 50, 14, 12, 30, 12],
    )

    row += 1
    row = write_section_title(ws, row, "??D6 ???????", span_cols=7)
    d6_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            d6_rows.append([i, action.get("action", "____"), action.get("target", "____"), action.get("responsible", "____"), action.get("due_date", "____"), "____(???/???/???)", "____"])
        else:
            d6_rows.append([i, str(action), "____", "____", "____", "____(???/???/???)", "____"])
    while len(d6_rows) < 3:
        d6_rows.append([len(d6_rows) + 1, "____", "____", "____", "____", "____", "____"])

    row = write_table(ws, row,
        ["??", "????", "????", "???", "????", "??", "????"],
        d6_rows,
        col_widths=[6, 40, 12, 14, 14, 18, 22],
    )

    # ==================== D7 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D7  ????? - ????(Prevent Recurrence - Yokoten)", span_cols=6)
    row += 1
    row = write_section_title(ws, row, "??????????", span_cols=5)
    d7 = template.get("d7_template", {})
    yokoten = d7.get("yokoten", [])
    y_rows = [[i + 1, item, "____(????/??/??)", "____", "____", "____"] for i, item in enumerate(yokoten)]
    while len(y_rows) < 4:
        y_rows.append([len(y_rows) + 1, "____(??? Yokoten ??)", "____", "____", "____", "____"])
    row = write_table(ws, row,
        ["??", "??????", "????", "???", "????", "??"],
        y_rows,
        col_widths=[6, 50, 22, 14, 14, 14],
    )

    row += 1
    row = write_section_title(ws, row, "??PFMEA ??", span_cols=2)
    row = write_kv_block(ws, row, [
        ("??????", "____"),
        ("? RPN", "____"),
        ("??? RPN", "____"),
        ("PFMEA ??? / ??", "____"),
    ], label_col_width=24, value_col_width=60)

    # ==================== D8 ====================
    row += 2
    row = write_yellow_section_title(ws, row, "D8  ???????(Team Recognition & Closure)", span_cols=4)
    row += 1
    row = write_section_title(ws, row, "??????", span_cols=2)
    row = write_kv_block(ws, row, [
        ("?? CA ??????", "____(?/?)"),
        ("?? CA ??????", "____(?/?)"),
        ("??????", "____(?/?,???????)"),
        ("????????(SOP/PFMEA/????)", "____(?/?)"),
        ("????????", "____(?/?)"),
        ("????", "____"),
        ("????", "____(???? / ????,??:____)"),
    ], label_col_width=34, value_col_width=60)

    row += 1
    row = write_section_title(ws, row, "??????", span_cols=2)
    row = write_kv_block(ws, row, [
        ("????????", "____"),
        ("?????", "____"),
        ("????????", "____"),
        ("?????????", "____"),
    ], label_col_width=24, value_col_width=60)

    row += 1
    row = write_section_title(ws, row, "?????", span_cols=4)
    row = write_table(ws, row,
        ["??", "??", "??", "??"],
        [
            ["??(?????)", "____", "____", "____"],
            ["??(????)", "____", "____", "____"],
            ["??(????)", "____", "____", "____"],
            ["????(??)", "____", "____", "____"],
        ],
        col_widths=[24, 20, 25, 18],
    )

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    cell = ws.cell(row=row, column=1, value=f"8D ????:{report_number}")
    cell.font = Font(name=FONT_NAME, size=10, italic=True, color="666666")
    cell.alignment = Alignment(horizontal="right", vertical="center")

    wb.save(output_path)
    print(f"[OK] Excel ???:{output_path}")


# ============================================================
# Word ??
# ============================================================

def set_cell_bg(cell, color_hex):
    """?? Word ?????????"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """?? Word ?????????"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, font_name=FONT_NAME, size=10.5, bold=False, color=None):
    """?? run ???(?? + ??)?"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # ????
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text, bold=False, size=10.5, color=None, alignment=None, indent=False):
    """?????"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    """??????(?????,???? Word ?? Heading ??????)?"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, color=HEADER_FILL)
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=SUBHEADER_FILL)
    else:
        set_run_font(run, size=11, bold=True, color="333333")
    return p


def add_table(doc, headers, rows, col_widths_cm=None, root_cause_row_indices=None):
    """
    ?????????
    headers: [str]
    rows: [[str]]
    col_widths_cm: [float]
    root_cause_row_indices: list of 0-based row indices that should be yellow-highlighted
    """
    if root_cause_row_indices is None:
        root_cause_row_indices = []

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # ????
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # ??
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        set_cell_bg(cell, HEADER_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, color="FFFFFF")

    # ???
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        is_rc = r_idx in root_cause_row_indices
        for c_idx, value in enumerate(row_data):
            cell = row_cells[c_idx]
            cell.text = ""
            if is_rc:
                set_cell_bg(cell, ROOT_CAUSE_FILL)
            elif r_idx % 2 == 1:
                set_cell_bg(cell, ALT_ROW_FILL)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=10, bold=is_rc)

    return table


def add_kv_table(doc, kv_pairs, label_width_cm=4.5, value_width_cm=12):
    """???????(2 ?)?"""
    table = doc.add_table(rows=len(kv_pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (label, value) in enumerate(kv_pairs):
        # ???
        label_cell = table.rows[i].cells[0]
        label_cell.width = Cm(label_width_cm)
        label_cell.text = ""
        set_cell_bg(label_cell, SUBHEADER_FILL)
        set_cell_borders(label_cell)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, size=10, bold=True, color="FFFFFF")

        # ??
        value_cell = table.rows[i].cells[1]
        value_cell.width = Cm(value_width_cm)
        value_cell.text = ""
        if i % 2 == 1:
            set_cell_bg(value_cell, ALT_ROW_FILL)
        set_cell_borders(value_cell)
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = value_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(value))
        set_run_font(run, size=10)

    return table


def set_doc_default_font(doc):
    """??????????????"""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)


def set_page_header(doc, report_number):
    """????,???? 8D-???"""
    section = doc.sections[0]
    header = section.header
    # ??????
    for p in header.paragraphs:
        p.text = ""
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # report_number ??? 8D- ??,??????
    run = p.add_run(report_number)
    set_run_font(run, size=9, color="666666")


def generate_word(context, template, output_path, report_number):
    """?? Word ???"""
    doc = Document()

    # ??????
    set_doc_default_font(doc)

    # ????:A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # ??
    set_page_header(doc, report_number)

    # ???
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("8D ??????")
    set_run_font(title_run, size=20, bold=True, color=HEADER_FILL)

    # ????
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    sub_run = sub_p.add_run(f"????:{report_number}    ????:{datetime.datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(sub_run, size=10, color="666666")

    # ============ ??D0-D2 ?????? ============
    add_heading(doc, "??D0-D2 ???????????", level=1)

    d0_d2 = template.get("d0_d2", {})

    add_heading(doc, "1.1 ??????", level=2)
    add_kv_table(doc, [
        ("8D ????", report_number),
        ("??????", datetime.datetime.now().strftime("%Y-%m-%d")),
        ("?????", "____"),
        ("????", "???"),
    ])

    add_heading(doc, "1.2 ????", level=2)
    add_kv_table(doc, [
        ("????", context.get("customer", "____")),
        ("?????", "____"),
        ("??????", "____"),
        ("??????", "____"),
        ("??????", "____(?? / ?? / 8D ??? / ??)"),
    ])

    add_heading(doc, "1.3 ????", level=2)
    add_kv_table(doc, [
        ("????", context.get("product", "____")),
        ("???? / ???", "____"),
        ("???", "____"),
        ("????", context.get("batch_size", "____")),
        ("????", "____"),
        ("???", context.get("defect_rate", "____")),
        ("????", d0_d2.get("defect_level_hint", "____")),
        ("????", d0_d2.get("discovery_location_hint", "____")),
        ("??", d0_d2.get("impact_hint", "____")),
    ])

    add_heading(doc, "1.4 D2 ????(5W2H)", level=2)
    add_kv_table(doc, [
        ("What(????)", context.get("defect", "____")),
        ("When(????)", "____(?? / ??)"),
        ("Where(????)", "____(?? / ???? / ????)"),
        ("Who(????)", "____"),
        ("Why(??????)", "____(????? / ????)"),
        ("How(????)", "____(???? / ????)"),
        ("How many(????/?)", f"{context.get('defect_rate', '____')},?? {context.get('batch_size', '____')}"),
    ])

    add_heading(doc, "1.5 ????", level=2)
    problem_stmt = (
        f"???{context.get('product', '____')}?? {d0_d2.get('discovery_location_hint', '____')} "
        f"???{context.get('defect', '____')}???,"
        f"??? {context.get('defect_rate', '____')},"
        f"?????? {context.get('batch_size', '____')}?"
        f"?????? {d0_d2.get('impact_hint', '____')},?????? 8D ???????"
    )
    add_paragraph(doc, problem_stmt)

    add_heading(doc, "1.6 D1 ?????", level=2)
    add_table(
        doc,
        headers=["??", "??", "??", "????"],
        rows=[
            ["????(?????)", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["????", "____", "???", "____"],
            ["?????", "____", "???", "____"],
            ["SQE(????????)", "____", "???", "____"],
            ["????(??)", "____", "???", "____"],
        ],
        col_widths_cm=[5.5, 3.5, 2.5, 4.5],
    )

    # ============ ??D3 ?????? ============
    add_heading(doc, "??D3 ??????", level=1)

    add_heading(doc, "2.1 ??????", level=2)
    d3 = template.get("d3_template", {})
    actions = d3.get("containment_actions", [])
    while len(actions) < 5:
        actions.append("____(???????)")

    d3_rows = []
    for i, action in enumerate(actions[:5], start=1):
        d3_rows.append([str(i), action, "____", "____", "____(?100%????/?????)", "???"])

    add_table(
        doc,
        headers=["??", "????", "???", "????", "????", "??"],
        rows=d3_rows,
        col_widths_cm=[1.2, 6.5, 1.8, 2.0, 4.0, 1.5],
    )

    add_heading(doc, "2.2 ???????", level=2)
    add_kv_table(doc, [
        ("??????", context.get("defect_rate", "____")),
        ("??????", "____"),
        ("??????(24h?)", "____"),
        ("??????(72h?)", "____"),
        ("????", "____(?? / ???,????????)"),
        ("????????", "____"),
        ("??????", "____"),
    ])

    # ============ ??D4 ?????? ============
    add_heading(doc, "??D4 ??????", level=1)

    add_heading(doc, "3.1 5Why ??", level=2)
    d4 = template.get("d4_template", {})
    five_why = d4.get("5why_path", {})

    # Use new steps array format, fall back to old format
    steps = five_why.get("steps", [])
    why_rows = []
    root_cause_indices = []

    if steps:
        for idx, step in enumerate(steps):
            level = step.get("level", "____")
            question = step.get("question", "____")
            answer = step.get("answer", "____")
            evidence = step.get("evidence", "____")
            why_rows.append([level, question, answer, evidence])
            if "??" in str(level):
                root_cause_indices.append(idx)
    else:
        why1 = five_why.get("why1", "____")
        why2_hints = five_why.get("why2_hints", [])
        why3_hints = five_why.get("why3_hints", [])
        why4_hints = five_why.get("why4_hints", [])
        why5_root = five_why.get("why5_root", "____")
        why_rows = [
            ["??", five_why.get("problem", "____"), "-", "-"],
            ["Why 1", "?????????", why1, "____(??/??/??)"],
            ["Why 2", f"???:{why1[:30]}...", "\n".join(f". {h}" for h in why2_hints) if why2_hints else "____", "____"],
            ["Why 3", "???????????", "\n".join(f". {h}" for h in why3_hints) if why3_hints else "____", "____"],
            ["Why 4", "?????/???????", "\n".join(f". {h}" for h in why4_hints) if why4_hints else "____", "____"],
            ["Why 5(??)", "?????/???????", why5_root, "____(??????)"],
        ]
        root_cause_indices = [5]

    add_table(
        doc,
        headers=["??", "??", "??", "??"],
        rows=why_rows,
        col_widths_cm=[2.5, 4.5, 6.0, 3.0],
        root_cause_row_indices=root_cause_indices,
    )

    add_heading(doc, "3.2 ??? 6M ??", level=2)
    six_m = d4.get("6m_analysis", {})

    if isinstance(six_m, list):
        m_rows = []
        for item in six_m:
            m_name = item.get("m", "____")
            finding = item.get("finding", "____")
            judgment = item.get("judgment", "____")
            m_rows.append([m_name, finding, judgment])
        add_table(
            doc,
            headers=["6M ??", "????", "??"],
            rows=m_rows,
            col_widths_cm=[3.0, 8.5, 3.5],
        )
    else:
        m_rows = [
            ["Man(?)", six_m.get("man", "____"), "____", "____(?/?)"],
            ["Machine(?)", six_m.get("machine", "____"), "____", "____(?/?)"],
            ["Material(?)", six_m.get("material", "____"), "____", "____(?/?)"],
            ["Method(?)", six_m.get("method", "____"), "____", "____(?/?)"],
            ["Measurement(?)", six_m.get("measurement", "____"), "____", "____(?/?)"],
            ["Environment(?)", six_m.get("environment", "____"), "____", "____(?/?)"],
        ]
        add_table(
            doc,
            headers=["6M ??", "????", "??", "????"],
            rows=m_rows,
            col_widths_cm=[2.8, 7.0, 3.5, 2.2],
        )

    add_heading(doc, "3.3 ??????", level=2)
    rc_summary = d4.get("root_cause_summary", [])
    if rc_summary:
        rc_kv = []
        for rc in rc_summary:
            rc_id = rc.get("id", "____")
            rc_desc = rc.get("description", "____")
            rc_type = rc.get("type", "____")
            rc_kv.append((f"{rc_id}({rc_type})", rc_desc))
        add_kv_table(doc, rc_kv, label_width_cm=4.5, value_width_cm=12)
    else:
        add_kv_table(doc, [
            ("RC1(????)", "____(?? 5Why ? 1-2 ???)"),
            ("RC2(????)", "____(?? 5Why ? 3-4 ???)"),
            ("RC3(????)", "____(?? 5Why ? 5 ???)"),
        ], label_width_cm=4.5, value_width_cm=12)

    add_heading(doc, "3.4 ????????", level=2)
    verify_text = d4.get("verification", "")
    if verify_text:
        add_paragraph(doc, verify_text)
    else:
        add_kv_table(doc, [
            ("????", "____(???? / ???? / ????)"),
            ("????", "____"),
            ("????", "____(??? / ??????)"),
            ("??? / ??", "____"),
        ], label_width_cm=4.5, value_width_cm=12)

    # ============ ??D5-D6 ?????? ============
    add_heading(doc, "??D5-D6 ??????", level=1)

    add_heading(doc, "4.1 D5 CA ??????", level=2)
    d5_d6 = template.get("d5_d6_template", {})
    permanent_actions = d5_d6.get("permanent_actions", [])

    ca_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            action_text = action.get("action", "____")
            target = action.get("target", "____")
        else:
            action_text = str(action)
            target = "____"
        ca_rows.append([str(i), action_text, target, "____(?/?/?)", "____(???????)", "____(??/??)"])

    while len(ca_rows) < 3:
        ca_rows.append([str(len(ca_rows) + 1), "____(??? CA ??)", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["??", "CA ??", "????", "???", "????", "??"],
        rows=ca_rows,
        col_widths_cm=[1.0, 6.5, 1.8, 1.8, 4.0, 1.8],
    )

    add_heading(doc, "4.2 D6 ???????", level=2)
    d6_rows = []
    for i, action in enumerate(permanent_actions, start=1):
        if isinstance(action, dict):
            action_text = action.get("action", "____")
            target = action.get("target", "____")
            responsible = action.get("responsible", "____")
            due = action.get("due_date", "____")
        else:
            action_text = str(action)
            target = "____"
            responsible = "____"
            due = "____"
        d6_rows.append([str(i), action_text, target, responsible, due, "____", "____"])

    while len(d6_rows) < 3:
        d6_rows.append([str(len(d6_rows) + 1), "____", "____", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["??", "????", "????", "???", "????", "??", "????"],
        rows=d6_rows,
        col_widths_cm=[1.0, 5.5, 1.8, 1.8, 1.8, 2.5, 2.6],
    )

    # ============ ??D7 ????? ============
    add_heading(doc, "??D7 ?????", level=1)

    add_heading(doc, "5.1 ????(Yokoten)????", level=2)
    d7 = template.get("d7_template", {})
    yokoten = d7.get("yokoten", [])
    yokoten_rows = []
    for i, item in enumerate(yokoten, start=1):
        yokoten_rows.append([str(i), item, "____(????/??/??)", "____", "____", "____"])

    while len(yokoten_rows) < 4:
        yokoten_rows.append([str(len(yokoten_rows) + 1), "____(??? Yokoten ??)", "____", "____", "____", "____"])

    add_table(
        doc,
        headers=["??", "??????", "????", "???", "????", "??"],
        rows=yokoten_rows,
        col_widths_cm=[1.0, 6.5, 3.0, 1.8, 1.8, 1.8],
    )

    add_heading(doc, "5.2 PFMEA ??", level=2)
    add_kv_table(doc, [
        ("??????", "____"),
        ("? RPN", "____"),
        ("??? RPN", "____"),
        ("PFMEA ??? / ??", "____"),
    ], label_width_cm=4.5, value_width_cm=12)

    # ============ ??D8 ????? ============
    add_heading(doc, "??D8 ?????", level=1)

    add_heading(doc, "6.1 ????", level=2)
    add_kv_table(doc, [
        ("?? CA ??????", "____(?/?)"),
        ("?? CA ??????", "____(?/?)"),
        ("??????", "____(?/?,???????)"),
        ("????????", "____(?/?,SOP/PFMEA/????)"),
        ("????????", "____(?/?)"),
        ("????", "____"),
        ("????", "____(???? / ????,??:____)"),
    ], label_width_cm=6.0, value_width_cm=10.5)

    add_heading(doc, "6.2 ????", level=2)
    add_kv_table(doc, [
        ("????????", "____"),
        ("?????", "____"),
        ("????????", "____"),
        ("?????????", "____"),
    ], label_width_cm=4.5, value_width_cm=12)

    add_heading(doc, "6.3 ???", level=2)
    add_table(
        doc,
        headers=["??", "??", "??", "??"],
        rows=[
            ["??(?????)", "____", "____", "____"],
            ["??(????)", "____", "____", "____"],
            ["??(????)", "____", "____", "____"],
            ["????(??)", "____", "____", "____"],
        ],
        col_widths_cm=[5.0, 4.0, 4.5, 3.0],
    )

    # ??????
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    end_p.paragraph_format.space_before = Pt(12)
    end_run = end_p.add_run(f"8D ????:{report_number}")
    set_run_font(end_run, size=9, color="666666")

    doc.save(output_path)
    print(f"[OK] Word ???:{output_path}")


# ============================================================
# ???
# ============================================================

def parse_args():
    parser = argparse.ArgumentParser(
        description="8D ?????????(Excel + Word)"
    )
    parser.add_argument("--product", required=True, help="????")
    parser.add_argument("--defect", required=True, help="????")
    parser.add_argument("--customer", required=False, default="____", help="????")
    parser.add_argument("--defect-rate", required=False, default="____", help="???")
    parser.add_argument("--batch-size", required=False, default="____", help="????")
    parser.add_argument(
        "--template",
        required=False,
        default="generic-defect",
        choices=["paint-defect", "assembly-defect", "welding-defect", "dimensional-defect", "generic-defect"],
        help="???? slug",
    )
    parser.add_argument(
        "--output-dir",
        required=False,
        default=os.path.expanduser("~/Desktop"),
        help="????(?? ~/Desktop)",
    )
    return parser.parse_args()


def main():
    # Windows ?? UTF-8 ????(?? print ?????? Agent ????)
    import io
    if sys.platform == "win32":
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

    args = parse_args()

    # ?? context
    context = {
        "product": args.product,
        "defect": args.defect,
        "customer": args.customer,
        "defect_rate": args.defect_rate,
        "batch_size": args.batch_size,
    }

    # ?? 8D ??
    report_number = generate_8d_number()
    print(f"[INFO] 8D ????:{report_number}")

    # ????
    template = load_template(args.template, context)
    print(f"[INFO] ????:{template.get('slug', args.template)} - {template.get('name', '')}")

    # ????????
    output_dir = os.path.expanduser(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)
    print(f"[INFO] ????:{output_dir}")

    # ?????
    safe_product = safe_filename(args.product)
    safe_defect = safe_filename(args.defect)

    excel_filename = f"8D??_{safe_product}_{safe_defect}_??.xlsx"
    word_filename = f"8D??_{safe_product}_{safe_defect}.docx"

    excel_path = os.path.join(output_dir, excel_filename)
    word_path = os.path.join(output_dir, word_filename)

    # ?? Excel
    generate_excel(context, template, excel_path, report_number)

    # ?? Word
    generate_word(context, template, word_path, report_number)

    # ?? JSON ????????
    result = {
        "status": "success",
        "report_number": report_number,
        "template_slug": template.get("slug", args.template),
        "template_name": template.get("name", ""),
        "excel_path": excel_path,
        "word_path": word_path,
        "output_dir": output_dir,
    }
    print("\n[RESULT_JSON]")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
